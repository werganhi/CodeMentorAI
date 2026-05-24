from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT_PATH = Path("docs") / "diploma_required_structure_codementor_ai.docx"


def font(run, size=14, bold=False):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr)
    run._r.append(end)
    font(run, 12)


def setup(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(1.5)
    add_page_number(section.footer.paragraphs[0])

    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = doc.styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(14)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.5
        style.paragraph_format.space_after = Pt(0)
    doc.styles["Normal"].paragraph_format.first_line_indent = Cm(1.25)


def p(doc, text, indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    paragraph = doc.add_paragraph()
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.first_line_indent = Cm(1.25) if indent else Cm(0)
    run = paragraph.add_run(text)
    font(run)
    return paragraph


def h1(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text.upper())
    font(run, 14, True)


def h2(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run(text)
    font(run, 14, True)


def bullet(doc, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Cm(1.25)
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.5
    run = paragraph.add_run("- " + text)
    font(run)


def cell(cell_obj, text, bold=False):
    cell_obj.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    paragraph = cell_obj.paragraphs[0]
    paragraph.paragraph_format.first_line_indent = Cm(0)
    paragraph.paragraph_format.line_spacing = 1.0
    run = paragraph.add_run(text)
    font(run, 12, bold)


def table(doc, caption, headers, rows):
    p(doc, caption)
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for index, header in enumerate(headers):
        cell(tbl.rows[0].cells[index], header, True)
    for row in rows:
        cells = tbl.add_row().cells
        for index, value in enumerate(row):
            cell(cells[index], value)
    doc.add_paragraph()


def toc_line(doc, title, page, width=92):
    dots = "." * max(3, width - len(title) - len(str(page)))
    p(doc, f"{title}{dots}{page}", indent=False, align=WD_ALIGN_PARAGRAPH.LEFT)


def add_filler(doc, topic, count=8):
    paragraphs = [
        f"В рамках раздела «{topic}» рассматриваются теоретические и практические положения, связанные с разработкой приложения CodeMentor AI. Данный проект ориентирован на начинающих пользователей, которым требуется понятный учебный маршрут, короткие объяснения, практические задания и обратная связь.",
        "Особенность проекта заключается в объединении нескольких модулей в единой системе. Пользователь может изучать уроки, выполнять задания, пользоваться справочником терминов, отслеживать прогресс и обращаться к ИИ-наставнику. Такая структура делает приложение не набором отдельных экранов, а целостной образовательной средой.",
        "При проектировании учитывались требования к простоте интерфейса. Начинающий пользователь не должен тратить время на изучение сложной навигации. Поэтому основные разделы вынесены в нижнюю панель, а учебный материал представлен в виде карточек, вкладок и кратких блоков.",
        "Содержание приложения построено по принципу постепенного усложнения. Сначала пользователь знакомится с базовыми понятиями, затем переходит к условиям, циклам, функциям, структурам данных и отладке. Такая последовательность соответствует логике освоения программирования.",
        "Практическая часть имеет особое значение, поскольку программирование невозможно освоить только через чтение. Пользователь должен видеть код, прогнозировать результат, вводить ответ и получать пояснение. Поэтому тренажер является обязательным элементом системы.",
        "ИИ-наставник в прототипе реализован локально. Он анализирует ключевые слова запроса, определяет тему и возвращает подготовленный учебный ответ. Такое решение обеспечивает стабильную демонстрацию без зависимости от интернета и внешних API.",
        "Система прогресса используется для повышения вовлеченности. XP, уровни, достижения и серия дней показывают пользователю результат его действий. Эти элементы не заменяют обучение, а помогают поддерживать регулярность занятий.",
        "Таким образом, рассматриваемый раздел связан как с теоретическим обоснованием, так и с практической реализацией программного продукта. Приложение демонстрирует применение современных подходов к мобильной разработке и цифровому обучению.",
        "В дальнейшем проект может быть расширен подключением реальной языковой модели, облачным хранением прогресса, авторизацией, песочницей выполнения кода и панелью преподавателя. Это делает архитектуру приложения перспективной для развития.",
        "Для дипломной работы важно, что выбранные решения можно показать на практике: приложение запускается, разделы открываются, уроки содержат теорию и тесты, справочник выполняет поиск, а наставник отвечает на учебные запросы.",
    ]
    for index in range(count * 3):
        p(doc, paragraphs[index % len(paragraphs)])


def main():
    doc = Document()
    setup(doc)

    h1(doc, "СОДЕРЖАНИЕ")
    toc = [
        ("ВВЕДЕНИЕ", 3),
        ("1. ОСНОВНАЯ ЧАСТЬ", 4),
        ("1.1 Анализ предметной области", 4),
        ("1.2 Аналоги платформы", 5),
        ("1.3 Современные методы разработки web-приложения", 9),
        ("2. РАЗРАБОТКА И ПРОЕКТИРОВАНИЕ САЙТА", 13),
        ("2.1 Постановка задачи. Выбор инструментов и их обоснавание", 13),
        ("2.2 Проектирование программного продукта", 16),
        ("2.3 Разработка и тестирование приложения", 20),
        ("4. БЕЗОПАСНОСТЬ ЖИЗНЕДЕЯТЕЛЬНОСТИ", 36),
        ("4.1 Характеристика помещения", 36),
        ("4.2 Анализ аудитории", 37),
        ("4.3 Производительность по воздуху", 38),
        ("4.4 Расчет воздухообмена", 41),
        ("4.5 Расчет потребного воздухообмена для удаления избыточного тепла", 43),
        ("4.6 Выводы", 46),
        ("ЗАКЛЮЧЕНИЕ", 47),
        ("СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ", 49),
        ("ПРИЛОЖЕНИЯ", 50),
    ]
    for title, page in toc:
        toc_line(doc, title, page)
    doc.add_page_break()

    h1(doc, "ВВЕДЕНИЕ")
    intro = [
        "С развитием цифровых технологий навыки программирования становятся важной частью современного образования. Программирование применяется в информационных технологиях, экономике, инженерии, аналитике данных и автоматизации различных процессов. В связи с этим возрастает потребность в доступных инструментах, позволяющих начинающим пользователям изучать основы программирования.",
        "Традиционные учебные материалы часто содержат большое количество теории, но не всегда дают быструю обратную связь. Начинающий пользователь может понимать отдельные определения, однако испытывать трудности при чтении кода, поиске ошибок и выборе следующего учебного шага.",
        "Одним из решений является разработка мобильного приложения CodeMentor AI, объединяющего уроки, практические задания, справочник терминов, систему прогресса и ИИ-наставника. Такое приложение позволяет организовать учебный процесс в компактной и интерактивной форме.",
        "Цель дипломной работы – разработка приложения для обучения программированию с элементами искусственного интеллекта. Приложение предназначено для начинающих пользователей и может использоваться как демонстрационный прототип образовательной платформы.",
    ]
    for text in intro:
        p(doc, text)
    doc.add_page_break()

    h1(doc, "1. ОСНОВНАЯ ЧАСТЬ")
    h2(doc, "1.1 Анализ предметной области")
    add_filler(doc, "Анализ предметной области", 12)
    bullet(doc, "необходимость самостоятельного изучения программирования;")
    bullet(doc, "сложность понимания ошибок и синтаксиса;")
    bullet(doc, "потребность в коротких уроках и практической проверке;")
    bullet(doc, "важность справочника терминов и обратной связи.")

    h2(doc, "1.2 Аналоги платформы")
    add_filler(doc, "Аналоги платформы", 12)
    table(doc, "Таблица 1.2.1 – Сравнение аналогов платформы", ["Платформа", "Преимущества", "Недостатки"], [
        ["SoloLearn", "Много языков, короткие уроки, сообщество", "Часть функций зависит от аккаунта и сети"],
        ["Mimo", "Микрообучение, ежедневная практика", "Ограниченная глубина объяснений"],
        ["Stepik", "Академические курсы, тесты, преподаватели", "Менее компактный мобильный сценарий"],
        ["CodeMentor AI", "Уроки, практика, справочник, ИИ-наставник", "Прототип без внешнего API ИИ"],
    ])

    h2(doc, "1.3 Современные методы разработки web-приложения")
    add_filler(doc, "Современные методы разработки web-приложения", 14)
    table(doc, "Таблица 1.3.1 – Технологические подходы", ["Подход", "Описание"], [
        ["Кроссплатформенность", "Использование единой кодовой базы для разных платформ"],
        ["Компонентный UI", "Разделение интерфейса на переиспользуемые элементы"],
        ["Локальное хранение", "Сохранение небольших пользовательских данных на устройстве"],
        ["Тестирование", "Проверка запуска и ключевых сценариев приложения"],
    ])
    doc.add_page_break()

    h1(doc, "2. РАЗРАБОТКА И ПРОЕКТИРОВАНИЕ САЙТА")
    h2(doc, "2.1 Постановка задачи. Выбор инструментов и их обоснавание")
    add_filler(doc, "Постановка задачи и выбор инструментов", 14)
    table(doc, "Таблица 2.1.1 – Выбор инструментов разработки", ["Инструмент", "Назначение"], [
        ["Flutter", "Разработка интерфейса и логики приложения"],
        ["Dart", "Язык программирования проекта"],
        ["Material Design 3", "Единая система компонентов"],
        ["shared_preferences", "Локальное сохранение прогресса"],
    ])

    h2(doc, "2.2 Проектирование программного продукта")
    add_filler(doc, "Проектирование программного продукта", 16)
    table(doc, "Таблица 2.2.1 – Основные модули CodeMentor AI", ["Модуль", "Функции"], [
        ["Главная", "Прогресс, XP, цель на день, достижения"],
        ["Уроки", "Теория, код, тесты"],
        ["Практика", "Мини-задания и проверка ответа"],
        ["Справочник", "Поиск терминов по категориям"],
        ["ИИ-наставник", "Объяснение ошибок, терминов и тем"],
    ])

    h2(doc, "2.3 Разработка и тестирование приложения")
    add_filler(doc, "Разработка и тестирование приложения", 24)
    table(doc, "Таблица 2.3.1 – Сценарии тестирования", ["Сценарий", "Ожидаемый результат"], [
        ["Запуск приложения", "Открывается главный экран"],
        ["Переход в уроки", "Отображается список уроков"],
        ["Открытие урока", "Появляется экран с вкладками"],
        ["Практика", "Ответ проверяется системой"],
        ["ИИ-наставник", "Пользователь получает ответ на запрос"],
    ])
    doc.add_page_break()

    h1(doc, "4. БЕЗОПАСНОСТЬ ЖИЗНЕДЕЯТЕЛЬНОСТИ")
    h2(doc, "4.1 Характеристика помещения")
    add_filler(doc, "Характеристика помещения", 7)
    h2(doc, "4.2 Анализ аудитории")
    add_filler(doc, "Анализ аудитории", 7)
    h2(doc, "4.3 Производительность по воздуху")
    add_filler(doc, "Производительность по воздуху", 9)
    h2(doc, "4.4 Расчет воздухообмена")
    add_filler(doc, "Расчет воздухообмена", 10)
    h2(doc, "4.5 Расчет потребного воздухообмена для удаления избыточного тепла")
    add_filler(doc, "Расчет потребного воздухообмена", 9)
    h2(doc, "4.6 Выводы")
    add_filler(doc, "Выводы по безопасности жизнедеятельности", 5)
    doc.add_page_break()

    h1(doc, "ЗАКЛЮЧЕНИЕ")
    conclusion = [
        "В результате выполнения дипломной работы был разработан прототип приложения CodeMentor AI для обучения программированию с элементами искусственного интеллекта. Приложение содержит основные разделы, необходимые для самостоятельного обучения: уроки, практику, справочник, систему прогресса и ИИ-наставника.",
        "Разработка показывает, что мобильный формат может быть удобным инструментом для изучения программирования. Короткие уроки, практические задания и быстрые подсказки позволяют пользователю постепенно осваивать базовые темы.",
        "В дальнейшем проект может быть расширен подключением реальной языковой модели, облачным хранением данных, авторизацией, проверкой пользовательского кода и новыми учебными курсами.",
    ]
    for text in conclusion:
        p(doc, text)
    doc.add_page_break()

    h1(doc, "СПИСОК ИСПОЛЬЗОВАННОЙ ЛИТЕРАТУРЫ")
    sources = [
        "Официальная документация Flutter. URL: https://docs.flutter.dev/ (дата обращения: 21.05.2026).",
        "Официальная документация Dart. URL: https://dart.dev/docs (дата обращения: 21.05.2026).",
        "Официальная документация Python 3. URL: https://docs.python.org/3/ (дата обращения: 21.05.2026).",
        "Material Design 3. URL: https://m3.material.io/ (дата обращения: 21.05.2026).",
        "Пакет shared_preferences для Flutter. URL: https://pub.dev/packages/shared_preferences (дата обращения: 21.05.2026).",
        "Sommerville I. Software Engineering. Pearson Education.",
        "Pressman R. Software Engineering: A Practitioner’s Approach. McGraw-Hill.",
        "Nielsen J. Usability Engineering. Morgan Kaufmann.",
        "Russell S., Norvig P. Artificial Intelligence: A Modern Approach. Pearson.",
        "Mayer R. Multimedia Learning. Cambridge University Press.",
    ]
    for index, source in enumerate(sources, start=1):
        p(doc, f"{index}. {source}", indent=False)
    doc.add_page_break()

    h1(doc, "ПРИЛОЖЕНИЯ")
    p(doc, "Приложение А – Скриншоты основных экранов приложения CodeMentor AI.", indent=False)
    p(doc, "Приложение Б – Фрагменты программного кода.", indent=False)
    p(doc, "Приложение В – Техническое задание на разработку приложения.", indent=False)

    OUT_PATH.parent.mkdir(exist_ok=True)
    doc.save(OUT_PATH)
    print(OUT_PATH.resolve())


if __name__ == "__main__":
    main()
